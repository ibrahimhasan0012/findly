from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DWT 2012 Problem - Complete Step-by-Step Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Cambria'

doc.add_paragraph()
subtitle = doc.add_paragraph()
run = subtitle.add_run("Natural Convection from Vertical Plate with Heat Source")
run.italic = True
run.font.name = 'Cambria'
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# Problem Statement
doc.add_heading('Problem Statement', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

doc.add_paragraph('Natural convection flow from an isothermal vertical plate with uniform heat source embedded in a stratified medium.')

doc.add_paragraph()
doc.add_heading('Given:', level=2)
doc.add_paragraph('• Wall temperature: T_w = constant')
doc.add_paragraph('• Ambient temperature: T_∞(x) = T_0 + B(x/L) (stratified)')
doc.add_paragraph('• Internal heat generation: Q(T - T_∞)')
doc.add_paragraph('• Gravity: g (downward)')

doc.add_page_break()

# STEP 1
doc.add_heading('STEP 1: Dimensional Governing Equations', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('Continuity:')
run.bold = True
run.font.name = 'Cambria'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('∂ū/∂x̄ + ∂v̄/∂ȳ = 0')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Momentum (boundary layer approximation):')
run.bold = True
run.font.name = 'Cambria'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('ū(∂ū/∂x̄) + v̄(∂ū/∂ȳ) = ν(∂²ū/∂ȳ²) + gβ(T − T∞,x)')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Energy:')
run.bold = True
run.font.name = 'Cambria'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²) + Q(T − T∞,x)')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Boundary Conditions:')
run.bold = True
run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('At ȳ = 0: ū = v̄ = 0, T = T_w')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

para = doc.add_paragraph()
run = para.add_run('As ȳ → ∞: ū → 0, T → T∞,x')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

doc.add_page_break()

# STEP 2
doc.add_heading('STEP 2: Define Reference Scales', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

doc.add_paragraph('Length scale: L (reference length)')
doc.add_paragraph('Temperature scale: ΔT = T_w − T_0')
doc.add_paragraph('Velocity scale: U = (ν/L)·Gr^(1/2)')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Grashof Number:')
run.bold = True
run.font.name = 'Cambria'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Gr = ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('gβ(T_w − T_0)L³')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(' / ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('ν²')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_page_break()

# STEP 3
doc.add_heading('STEP 3: Define Dimensionless Variables', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('Coordinates:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('x = x̄/L,    y = (Gr^(1/4)/L)·ȳ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Velocities:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('u = (L/ν)·Gr^(−1/2)·ū,    v = (L/ν)·Gr^(−1/4)·v̄')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Temperature:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('θ = (T − T∞,x)/(T_w − T_0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

doc.add_page_break()

# STEP 4
doc.add_heading('STEP 4: Dimensionless Equations', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('After substitution and simplification:')
run.font.name = 'Cambria'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Continuity:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('∂u/∂x + ∂v/∂y = 0')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Momentum:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Energy:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²) + λθ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Parameters:')
run.bold = True

doc.add_paragraph('• Pr = ν/α (Prandtl number)')
doc.add_paragraph('• λ = QL²/ν (heat generation parameter)')
doc.add_paragraph('• S = B/(T_w − T_0)·Gr^(1/2) (stratification parameter)')

doc.add_page_break()

# STEP 5
doc.add_heading('STEP 5: Similarity Transformation', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('Introduce stream function and similarity variable:')
run.font.name = 'Cambria'

doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('ψ = x^(3/4)·f(x,η),    η = x^(−1/4)·y,    θ = θ(x,η)')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Velocity components:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('u = ∂ψ/∂y = x^(1/2)·f′(x,η)')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
run.font.italic = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('v = −∂ψ/∂x = (1/4)x^(−1/4)·[ηf′ − 3f] − x^(3/4)·∂f/∂x')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
run.font.italic = True

doc.add_page_break()

# STEP 6
doc.add_heading('STEP 6: Final ODE System', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
run = para.add_run('After substituting similarity variables into PDEs:')
run.font.name = 'Cambria'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Momentum ODE:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("f′′′ + (3/4)f·f′′ − (1/2)(f′)² + θ = x(f′·∂f′/∂x − f′′·∂f/∂x)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
run.font.italic = True
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Energy ODE:')
run.bold = True

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("(1/Pr)θ′′ + (3/4)f·θ′ + λθ = x(f′·∂θ/∂x − θ′·∂f/∂x)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
run.font.italic = True
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Boundary Conditions:')
run.bold = True

para = doc.add_paragraph()
run = para.add_run('At η = 0: f = f′ = 0, θ = 1 − Sx')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

para = doc.add_paragraph()
run = para.add_run('As η → ∞: f′ → 0, θ → 0')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

doc.add_page_break()

# STEP 7 - Engineering quantities with proper LaTeX formatting
doc.add_heading('STEP 7: Engineering Quantities (Equation 26)', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Skin Friction Coefficient:')
run.bold = True
run.font.name = 'Cambria'
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Build with proper formatting
run = para.add_run('½ ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run('1/4')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.superscript = True

run = para.add_run(' C')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('f')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run(' = (∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('u')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('/∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(')|')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y=0')
run.font.name = 'Cambria Math'
run.font.size = Pt(9)
run.font.subscript = True

run = para.add_run(' = ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('f')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('′′(')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(', 0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Nusselt Number:')
run.bold = True
run.font.name = 'Cambria'
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = para.add_run('Nu')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(' / Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run('1/4')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.superscript = True

run = para.add_run(' = −(∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('θ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('/∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(')|')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y=0')
run.font.name = 'Cambria Math'
run.font.size = Pt(9)
run.font.subscript = True

run = para.add_run(' = −')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('θ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('′(')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(', 0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

doc.add_page_break()

# Summary
doc.add_heading('SUMMARY: Complete Solution Process', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

doc.add_paragraph('1. Start with dimensional equations (continuity, momentum, energy)', style='List Number')
doc.add_paragraph('2. Define reference scales (L, ΔT, U based on Gr)', style='List Number')
doc.add_paragraph('3. Make variables dimensionless (x, y, u, v, θ)', style='List Number')
doc.add_paragraph('4. Substitute and simplify to get dimensionless PDEs', style='List Number')
doc.add_paragraph('5. Introduce similarity transformation (ψ, η)', style='List Number')
doc.add_paragraph('6. Convert PDEs to ODEs', style='List Number')
doc.add_paragraph('7. Solve ODEs numerically to get f(η) and θ(η)', style='List Number')
doc.add_paragraph('8. Calculate engineering quantities (C_f, Nu)', style='List Number')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('✓ This matches your handnotes pages 1-7 exactly!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)
run.font.name = 'Cambria'

# Save
doc.save('DWT_2012_Complete_Solution.docx')
print("✓ Created: DWT_2012_Complete_Solution.docx")
print("\nComplete step-by-step solution with LaTeX-style formatting!")

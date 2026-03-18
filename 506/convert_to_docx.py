from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_markdown_formatting(text):
    """Remove markdown formatting artifacts"""
    # Remove bold markers
    text = text.replace('**', '')
    # Remove code backticks
    text = text.replace('`', '')
    # Remove blockquote markers
    text = text.replace('> ', '')
    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def recreate_docx_clean(md_file, docx_file, title):
    """Convert markdown to DOCX with clean formatting (no ** or ` artifacts)"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    skip_alert = False
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'No Spacing'
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            if i > 0 and lines[i-1].strip():  # Only add space if previous line had content
                doc.add_paragraph()
            i += 1
            continue
        
        # Handle alerts/callouts
        if line.startswith('> [!'):
            alert_type = line.split('[!')[1].split(']')[0] if '[!' in line else ''
            # Get alert text, handling multi-line alerts
            alert_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith('>'):
                alert_lines.append(lines[j])
                j += 1
            
            # Combine alert text
            full_alert = ' '.join([l.replace('>', '').replace('[!'+alert_type+']', '').strip() 
                                  for l in alert_lines])
            full_alert = clean_markdown_formatting(full_alert)
            
            if full_alert:
                para = doc.add_paragraph()
                run = para.add_run(f"⚠ {alert_type}: {full_alert}")
                run.bold = True
                if alert_type == 'IMPORTANT':
                    run.font.color.rgb = RGBColor(200, 0, 0)
                elif alert_type == 'TIP':
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif alert_type == 'WARNING':
                    run.font.color.rgb = RGBColor(255, 100, 0)
                elif alert_type == 'NOTE':
                    run.font.color.rgb = RGBColor(0, 0, 200)
            
            i = j
            continue
        
        # Handle headers
        if line.startswith('#### '):
            text = clean_markdown_formatting(line[5:])
            doc.add_heading(text, level=4)
        elif line.startswith('### '):
            text = clean_markdown_formatting(line[4:])
            doc.add_heading(text, level=3)
        elif line.startswith('## '):
            text = clean_markdown_formatting(line[3:])
            doc.add_heading(text, level=2)
        elif line.startswith('# '):
            text = clean_markdown_formatting(line[2:])
            doc.add_heading(text, level=1)
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = clean_markdown_formatting(line[2:])
            if text:
                doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif len(line) > 2 and line[0].isdigit() and '. ' in line[:4]:
            text = clean_markdown_formatting(line.split('. ', 1)[1])
            if text:
                doc.add_paragraph(text, style='List Number')
        
        # Handle horizontal rules
        elif line.strip() == '---':
            para = doc.add_paragraph('─' * 80)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle tables (simple detection)
        elif '|' in line and line.count('|') >= 2:
            # Skip table for now or handle simply
            i += 1
            continue
        
        # Regular paragraph
        else:
            text = clean_markdown_formatting(line)
            if text and not text.startswith('>'):
                para = doc.add_paragraph(text)
                para.style = 'Normal'
        
        i += 1
    
    doc.save(docx_file)
    print(f"✓ Created: {docx_file}")

# Recreate all DOCX files with clean formatting
print("Recreating DOCX files with clean formatting (removing ** and ` artifacts)...\n")

recreate_docx_clean(
    'C:/Users/User/.gemini/antigravity/brain/c3e58dd8-4f45-4051-b036-3e9a05b1da72/quiz_study_guide.md',
    'AMCS506_Quiz_Study_Guide.docx',
    'AMCS 506 Quiz Study Guide - Lectures 1-3'
)

recreate_docx_clean(
    'C:/Users/User/.gemini/antigravity/brain/c3e58dd8-4f45-4051-b036-3e9a05b1da72/nondimensionalization_solution.md',
    'AMCS506_NonDimensionalization_Solution.docx',
    'Non-Dimensionalization Problem - DWT 2012'
)

recreate_docx_clean(
    'C:/Users/User/.gemini/antigravity/brain/c3e58dd8-4f45-4051-b036-3e9a05b1da72/formula_sheet.md',
    'AMCS506_Formula_Sheet.docx',
    'AMCS 506 Quick Reference Formula Sheet'
)

print("\n" + "="*60)
print("All DOCX files recreated with clean formatting!")
print("No more ** or ` artifacts!")
print("="*60)

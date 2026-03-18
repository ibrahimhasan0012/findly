from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('AMCS 506 - Complete Handnote Breakdown & Practice Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.add_run("Detailed explanations of every formula with practice strategies").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# ============= PAGE 1 BREAKDOWN =============
doc.add_heading('Page 1: Prandtl Boundary Layer Equations', level=1)

if os.path.exists("handnote_page_1.png"):
    doc.add_picture("handnote_page_1.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Formula 1: Continuity Equation', level=2)

para = doc.add_paragraph()
run = para.add_run('∂u/∂x + ∂v/∂y = 0')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why this formula?', level=3)
doc.add_paragraph('This is the mass conservation law for incompressible flow. It states that what flows in must flow out - mass cannot be created or destroyed.')

doc.add_heading('Physical meaning:', level=3)
doc.add_paragraph('If velocity increases in x-direction (∂u/∂x > 0), then velocity must decrease in y-direction (∂v/∂y < 0) to maintain constant mass flow.')

doc.add_heading('When to use:', level=3)
doc.add_paragraph('ALWAYS start with continuity when solving ANY fluid flow problem. It\'s the foundation.')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Given u = x² + y, find v such that continuity is satisfied', style='List Bullet')
doc.add_paragraph('Practice 2: Verify if u = xy, v = -xy satisfies continuity', style='List Bullet')
doc.add_paragraph('Practice 3: For u = 3x²y, derive v by integrating ∂v/∂y = -∂u/∂x', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Formula 2: Momentum Equation (Boundary Layer)', level=2)

para = doc.add_paragraph()
run = para.add_run('u(∂u/∂x) + v(∂u/∂y) = ν(∂²u/∂y²) + gβ(T - T∞)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why this formula?', level=3)
doc.add_paragraph('This is Newton\'s second law (F = ma) applied to fluid flow with buoyancy. The boundary layer approximation removes ∂²u/∂x² term.')

doc.add_heading('Term-by-term breakdown:', level=3)
doc.add_paragraph('u(∂u/∂x): Convective acceleration in x-direction (fluid speeding up as it moves)', style='List Bullet')
doc.add_paragraph('v(∂u/∂y): Convective acceleration in y-direction (fluid moving perpendicular)', style='List Bullet')
doc.add_paragraph('ν(∂²u/∂y²): Viscous diffusion (friction slowing down the fluid)', style='List Bullet')
doc.add_paragraph('gβ(T - T∞): Buoyancy force (hot fluid rises, cold sinks)', style='List Bullet')

doc.add_heading('Why only ∂²u/∂y²?', level=3)
doc.add_paragraph('Boundary layer is THIN. Changes perpendicular to wall (y-direction) are much larger than changes parallel to wall (x-direction). So ∂²u/∂x² << ∂²u/∂y² and we drop it!')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Identify each term\'s physical meaning in a given momentum equation', style='List Bullet')
doc.add_paragraph('Practice 2: Estimate order of magnitude: if δ = 0.01m and L = 1m, show ∂²u/∂y² >> ∂²u/∂x²', style='List Bullet')
doc.add_paragraph('Practice 3: Write momentum equation for forced convection (no buoyancy term)', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Formula 3: Energy Equation (Boundary Layer)', level=2)

para = doc.add_paragraph()
run = para.add_run('u(∂T/∂x) + v(∂T/∂y) = α(∂²T/∂y²)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
para.add_run('where α = k/(ρCₚ) = thermal diffusivity').italic = True

doc.add_heading('Why this formula?', level=3)
doc.add_paragraph('Energy conservation - heat is transported by convection (left side) and diffusion (right side).')

doc.add_heading('Physical meaning:', level=3)
doc.add_paragraph('Left side: Heat carried by moving fluid (like wind carrying heat)', style='List Bullet')
doc.add_paragraph('Right side: Heat spreading by conduction (like heat spreading through metal)', style='List Bullet')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: If Pr = ν/α = 0.7 (air), which is faster: momentum or heat diffusion?', style='List Bullet')
doc.add_paragraph('Practice 2: Write energy equation with heat source Q', style='List Bullet')
doc.add_paragraph('Practice 3: Compare thermal vs velocity boundary layer thickness using Pr', style='List Bullet')

doc.add_page_break()

# ============= PAGE 2 BREAKDOWN =============
doc.add_heading('Page 2: Non-Dimensionalization - Chain Rule Applications', level=1)

if os.path.exists("handnote_page_2.png"):
    doc.add_picture("handnote_page_2.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('The Chain Rule Technique', level=2)

para = doc.add_paragraph()
run = para.add_run('⚠ CRITICAL SKILL: This is the HEART of non-dimensionalization!')
run.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph()
doc.add_heading('Formula: Chain Rule for ∂u/∂x', level=2)

para = doc.add_paragraph()
run = para.add_run('∂ū/∂x̄ = (∂u/∂x̄) · (∂x/∂x̄) · (∂ū/∂u)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Step-by-step breakdown:', level=3)

doc.add_paragraph('Step 1: Identify what you have')
doc.add_paragraph('ū = dimensional velocity', style='List Bullet')
doc.add_paragraph('x̄ = dimensional coordinate', style='List Bullet')
doc.add_paragraph('u = dimensionless velocity = ū/U (where U is velocity scale)', style='List Bullet')
doc.add_paragraph('x = dimensionless coordinate = x̄/L (where L is length scale)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Step 2: Apply chain rule')
doc.add_paragraph('∂x/∂x̄ = 1/L (since x = x̄/L)', style='List Bullet')
doc.add_paragraph('∂ū/∂u = U (since ū = U·u)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Step 3: Substitute')
para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ = (∂u/∂x̄) · (1/L) · U = (U/L) · (∂u/∂x)').font.name = 'Courier New'

doc.add_heading('Why this matters:', level=3)
doc.add_paragraph('This converts dimensional derivatives to dimensionless ones. You MUST master this for the quiz!')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Derive ∂v̄/∂ȳ in terms of ∂v/∂y using chain rule', style='List Bullet')
doc.add_paragraph('Practice 2: Show ∂²ū/∂ȳ² = (U·Gr^(1/2)/L²)·(∂²u/∂y²) for boundary layer scaling', style='List Bullet')
doc.add_paragraph('Practice 3: Apply chain rule to ∂T/∂x where T = T∞ + ΔT·θ', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Formula: Stretched Coordinate', level=2)

para = doc.add_paragraph()
run = para.add_run('y = ȳ/L · Gr^(1/4)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why stretch y?', level=3)
doc.add_paragraph('Boundary layer thickness δ ~ L/Gr^(1/4). Without stretching, y would be very small (0.001, 0.002...). Stretching makes it O(1) so equations are easier to solve numerically.')

doc.add_heading('Physical analogy:', level=3)
doc.add_paragraph('Like zooming into a microscope - the boundary layer is thin, so we "zoom in" to see details.')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: If Gr = 10^8, what is the stretching factor? (Answer: 10^2)', style='List Bullet')
doc.add_paragraph('Practice 2: Show that ∂/∂ȳ = (Gr^(1/4)/L)·(∂/∂y)', style='List Bullet')
doc.add_paragraph('Practice 3: Explain why u and v have different scalings', style='List Bullet')

doc.add_page_break()

# ============= PAGE 3 BREAKDOWN =============
doc.add_heading('Page 3: Temperature Non-Dimensionalization', level=1)

if os.path.exists("handnote_page_3.png"):
    doc.add_picture("handnote_page_3.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Formula: Dimensionless Temperature', level=2)

para = doc.add_paragraph()
run = para.add_run('θ = (T - T∞)/ΔT')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why this definition?', level=3)
doc.add_paragraph('Makes temperature dimensionless and scaled between 0 and 1:')
doc.add_paragraph('θ = 0 at free stream (T = T∞)', style='List Bullet')
doc.add_paragraph('θ = 1 at wall (T = T_w, so ΔT = T_w - T∞)', style='List Bullet')

doc.add_heading('Alternative form:', level=3)
para = doc.add_paragraph()
para.add_run('T = T∞ + ΔT·θ').font.name = 'Courier New'
doc.add_paragraph('This is useful when taking derivatives!')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: If T_w = 100°C, T∞ = 20°C, and T = 60°C, find θ (Answer: 0.5)', style='List Bullet')
doc.add_paragraph('Practice 2: Derive ∂T/∂x in terms of θ using chain rule', style='List Bullet')
doc.add_paragraph('Practice 3: Show ∂²T/∂y² = (ΔT·Gr^(1/2)/L²)·(∂²θ/∂y²)', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Formula: Energy Equation Transformation', level=2)

para = doc.add_paragraph()
run = para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)·(∂²θ/∂y²)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Where does 1/Pr come from?', level=3)
doc.add_paragraph('From dimensional analysis:')
doc.add_paragraph('α/ν = 1/Pr (since Pr = ν/α)', style='List Bullet')
doc.add_paragraph('This ratio controls relative importance of thermal vs momentum diffusion', style='List Bullet')

doc.add_heading('Physical interpretation:', level=3)
doc.add_paragraph('Pr = 0.7 (air): Heat diffuses slightly faster than momentum', style='List Bullet')
doc.add_paragraph('Pr = 7 (water): Momentum diffuses faster than heat', style='List Bullet')
doc.add_paragraph('Pr = 0.01 (liquid metal): Heat diffuses MUCH faster than momentum', style='List Bullet')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Derive the energy equation from dimensional form step-by-step', style='List Bullet')
doc.add_paragraph('Practice 2: Add heat source term λθ to the equation', style='List Bullet')
doc.add_paragraph('Practice 3: Compare thermal boundary layer thickness for Pr = 0.7 vs Pr = 7', style='List Bullet')

doc.add_page_break()

# ============= PAGE 4 BREAKDOWN =============
doc.add_heading('Page 4: Similarity Transformation', level=1)

if os.path.exists("handnote_page_4.png"):
    doc.add_picture("handnote_page_4.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Formula: Stream Function', level=2)

para = doc.add_paragraph()
run = para.add_run('ψ = x^(3/4)·f(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('η = x^(-1/4)·y')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('What is a stream function?', level=3)
doc.add_paragraph('A mathematical trick that AUTOMATICALLY satisfies continuity! If we define:')
doc.add_paragraph('u = ∂ψ/∂y', style='List Bullet')
doc.add_paragraph('v = -∂ψ/∂x', style='List Bullet')
doc.add_paragraph('Then ∂u/∂x + ∂v/∂y = 0 is automatically satisfied!')

doc.add_heading('Why x^(3/4)?', level=3)
doc.add_paragraph('This comes from similarity analysis. For natural convection boundary layers, the solution has self-similar structure with this specific power.')

doc.add_heading('Why η = x^(-1/4)·y?', level=3)
doc.add_paragraph('This is the similarity variable. It combines x and y into a single variable, reducing PDEs to ODEs!')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Verify that u = ∂ψ/∂y and v = -∂ψ/∂x satisfy continuity', style='List Bullet')
doc.add_paragraph('Practice 2: Calculate u and v from ψ = x^(3/4)·f(η)', style='List Bullet')
doc.add_paragraph('Practice 3: Show ∂u/∂y = x^(1/2)·x^(-1/4)·f\' = x^(1/4)·f\'', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Formula: Boundary Conditions', level=2)

para = doc.add_paragraph()
run = para.add_run('At y = 0: u = v = 0, θ = 1 + A sin(πx)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('As y → ∞: u → 0, θ → 0')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why these conditions?', level=3)
doc.add_paragraph('u = v = 0 at y = 0: No-slip condition (fluid sticks to wall)', style='List Bullet')
doc.add_paragraph('θ = 1 + A sin(πx): Sinusoidal wall temperature variation', style='List Bullet')
doc.add_paragraph('u → 0 as y → ∞: Velocity decays to zero far from wall', style='List Bullet')
doc.add_paragraph('θ → 0 as y → ∞: Temperature approaches ambient', style='List Bullet')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Transform these BCs to similarity variables (f, f\', θ)', style='List Bullet')
doc.add_paragraph('Practice 2: What does A = 0.3 mean physically?', style='List Bullet')
doc.add_paragraph('Practice 3: Write BCs for uniform wall temperature (no sinusoidal variation)', style='List Bullet')

doc.add_page_break()

# ============= PAGES 5-7 BREAKDOWN =============
doc.add_heading('Pages 5-7: Final ODE Derivation', level=1)

doc.add_paragraph('These pages show the complete transformation from PDEs to ODEs using the similarity transformation.')

doc.add_heading('Key Formula: Final ODE', level=2)

para = doc.add_paragraph()
run = para.add_run('f\'\'\' + (3/4)f·f\'\' - (1/2)(f\')² + θ = x(f\'·∂f\'/∂x - f\'\'·∂f/∂x)')
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.bold = True

doc.add_heading('Term-by-term meaning:', level=3)
doc.add_paragraph('f\'\'\': Third derivative - viscous diffusion', style='List Bullet')
doc.add_paragraph('(3/4)f·f\'\': Convection term', style='List Bullet')
doc.add_paragraph('-(1/2)(f\')²: Inertia term', style='List Bullet')
doc.add_paragraph('θ: Buoyancy driving force', style='List Bullet')
doc.add_paragraph('Right side: Non-similar terms (depend on x)', style='List Bullet')

doc.add_heading('Why this is important:', level=3)
doc.add_paragraph('This ODE is what you actually SOLVE numerically! Much easier than solving the original PDEs.')

doc.add_heading('How to practice:', level=3)
doc.add_paragraph('Practice 1: Identify which terms come from which original PDE terms', style='List Bullet')
doc.add_paragraph('Practice 2: For similar solution (∂/∂x = 0), simplify the equation', style='List Bullet')
doc.add_paragraph('Practice 3: Write the energy equation in similarity form', style='List Bullet')

doc.add_page_break()

# ============= MASTER PRACTICE STRATEGY =============
doc.add_heading('MASTER PRACTICE STRATEGY', level=1)

para = doc.add_paragraph()
run = para.add_run('🎯 How to Master Non-Dimensionalization for the Quiz')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
doc.add_heading('Week Before Quiz: Daily Practice Plan', level=2)

doc.add_heading('Day 1-2: Master the Chain Rule', level=3)
doc.add_paragraph('Exercise 1: Given x = x̄/L, derive ∂/∂x̄ in terms of ∂/∂x', style='List Bullet')
doc.add_paragraph('Exercise 2: Given u = ū/U, derive ∂ū/∂x̄ in terms of ∂u/∂x', style='List Bullet')
doc.add_paragraph('Exercise 3: Combine both: derive ∂ū/∂x̄ completely', style='List Bullet')
doc.add_paragraph('Exercise 4: Do the same for ∂²ū/∂ȳ² with stretched coordinate', style='List Bullet')

doc.add_heading('Day 3-4: Practice Full Non-Dimensionalization', level=3)
doc.add_paragraph('Take the dimensional momentum equation and non-dimensionalize it step-by-step', style='List Bullet')
doc.add_paragraph('Take the dimensional energy equation and non-dimensionalize it step-by-step', style='List Bullet')
doc.add_paragraph('Identify all dimensionless parameters (Re, Pr, Gr)', style='List Bullet')

doc.add_heading('Day 5-6: Similarity Transformations', level=3)
doc.add_paragraph('Practice converting PDEs to ODEs using η = x^(-1/4)·y', style='List Bullet')
doc.add_paragraph('Practice stream function derivatives', style='List Bullet')
doc.add_paragraph('Transform boundary conditions to similarity form', style='List Bullet')

doc.add_heading('Day 7 (Day Before Quiz): Review', level=3)
doc.add_paragraph('Go through your handwritten notes page by page', style='List Bullet')
doc.add_paragraph('Redo the DWT 2012 problem from scratch', style='List Bullet')
doc.add_paragraph('Make sure you can write all formulas from memory', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Common Mistakes to Avoid', level=2)

para = doc.add_paragraph()
run = para.add_run('❌ Mistake 1: Forgetting to stretch y-coordinate')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
doc.add_paragraph('✓ Always use y = ȳ/L · Gr^(1/4) for boundary layers!')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('❌ Mistake 2: Wrong velocity scaling')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
doc.add_paragraph('✓ u and v scale differently! u ~ Gr^(1/2), v ~ Gr^(1/4)')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('❌ Mistake 3: Missing factors in chain rule')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
doc.add_paragraph('✓ Always include ALL three parts: ∂f/∂x̄ = (∂f/∂x)·(∂x/∂x̄)·(∂f̄/∂f)')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('❌ Mistake 4: Not simplifying after substitution')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
doc.add_paragraph('✓ After substituting, collect terms and cancel Gr factors!')

doc.add_page_break()

# ============= QUICK REFERENCE =============
doc.add_heading('QUICK REFERENCE FOR QUIZ', level=1)

doc.add_heading('Essential Formulas to Memorize', level=2)

doc.add_paragraph('1. Continuity: ∂u/∂x + ∂v/∂y = 0')
doc.add_paragraph('2. Momentum: u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ')
doc.add_paragraph('3. Energy: u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²)')
doc.add_paragraph('4. Grashof: Gr = gβΔTL³/ν²')
doc.add_paragraph('5. Prandtl: Pr = ν/α')
doc.add_paragraph('6. Stretched y: y = ȳ/L · Gr^(1/4)')
doc.add_paragraph('7. Stream function: u = ∂ψ/∂y, v = -∂ψ/∂x')
doc.add_paragraph('8. Similarity: η = x^(-1/4)·y')

doc.add_paragraph()
doc.add_heading('Chain Rule Template', level=2)

para = doc.add_paragraph()
para.add_run('For any derivative ∂f̄/∂x̄:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Step 1: ∂f̄/∂x̄ = (∂f/∂x̄) · (∂x/∂x̄) · (∂f̄/∂f)').font.name = 'Courier New'

para = doc.add_paragraph()
para.add_run('Step 2: Substitute ∂x/∂x̄ = 1/L and ∂f̄/∂f = F (scale)').font.name = 'Courier New'

para = doc.add_paragraph()
para.add_run('Step 3: Result = (F/L) · (∂f/∂x)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Final Tips', level=2)

para = doc.add_paragraph()
run = para.add_run('✓ Show ALL steps - professors want to see your process!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

para = doc.add_paragraph()
run = para.add_run('✓ Write down what each variable represents')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

para = doc.add_paragraph()
run = para.add_run('✓ Check dimensions - everything should be dimensionless at the end!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

para = doc.add_paragraph()
run = para.add_run('✓ Practice writing formulas from memory - speed matters!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
doc.add_paragraph("─" * 80)
doc.add_paragraph()

para = doc.add_paragraph()
para.add_run('Good luck on your quiz! 🚀').bold = True
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
doc.save('AMCS506_Complete_Handnote_Breakdown_with_Practice.docx')
print("✓ Created: AMCS506_Complete_Handnote_Breakdown_with_Practice.docx")
print("\nComplete breakdown with practice strategies created!")

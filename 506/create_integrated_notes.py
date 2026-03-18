from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('AMCS 506 - Complete Study Guide: Handnotes + Lecture 3 Materials', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.add_run("Integrated handwritten notes with lecture slide content").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# ============= LECTURE 3 OVERVIEW =============
doc.add_heading('Lecture 3: Boundary Layer Equations - Overview', level=1)

para = doc.add_paragraph()
run = para.add_run('From Lecture Slides:')
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph('Discovered by Ludwig Prandtl (1875-1953)')
doc.add_paragraph('In fluid dynamics, the boundary layer is a thin region adjacent to a solid surface where the effects of viscosity are significant.')

doc.add_paragraph()
doc.add_heading('Key Concept from Slides:', level=2)
doc.add_paragraph('In laminar flow, the fluid moves in parallel layers with minimal mixing between them. The study of laminar 2D convective boundary layer flow is essential for applications in engineering and environmental science.')

doc.add_paragraph()
doc.add_heading('Why Boundary Layer Theory Matters:', level=2)
doc.add_paragraph('For the boundary layer flow: The viscous term w.r.t. x derivative term is negligible due to low contribution (found by order of magnitude analysis). See Book: Convective Heat Transfer, p 63-64')

doc.add_page_break()

# ============= PAGE 1 WITH LECTURE CONTENT =============
doc.add_heading('Section 1: Prandtl Boundary Layer Equations', level=1)

para = doc.add_paragraph()
run = para.add_run('📖 From Lecture Slides + Your Handnotes')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

if os.path.exists("handnote_page_1.png"):
    doc.add_picture("handnote_page_1.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Lecture 3 Slide Content:', level=2)

doc.add_paragraph('Laminar boundary layer equations (from slides):')
doc.add_paragraph()

para = doc.add_paragraph()
para.add_run('Continuity:').bold = True
doc.add_paragraph('∂u/∂x + ∂v/∂y = 0', style='No Spacing').runs[0].font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Momentum (x-direction):').bold = True
doc.add_paragraph('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + gβ(T - T∞)', style='No Spacing').runs[0].font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Energy:').bold = True
doc.add_paragraph('u(∂T/∂x) + v(∂T/∂y) = (k/ρCₚ)(∂²T/∂y²)', style='No Spacing').runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('From Your Handnotes:', level=2)
doc.add_paragraph('α = k/(ρCₚ) - thermal diffusivity')

doc.add_paragraph()
doc.add_heading('Flow Regimes (from slides):', level=2)
doc.add_paragraph('(a) Laminar flow', style='List Bullet')
doc.add_paragraph('(b) Transitional flow', style='List Bullet')
doc.add_paragraph('(c) Turbulent flow', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Turbulent Flow Notes (from handnotes):', level=2)
doc.add_paragraph('More chaotic - Richard Feynman described turbulence as "the most important unsolved problem of classical physics"')

doc.add_paragraph()
doc.add_heading('Numerical Methods (from handnotes):', level=2)
doc.add_paragraph('RANS - Reynolds Averaged Navier-Stokes', style='List Bullet')
doc.add_paragraph('LES - Large-Eddy Simulation (turbulent)', style='List Bullet')
doc.add_paragraph('DNS - Direct Numerical Simulation (very sufficient)', style='List Bullet')

doc.add_page_break()

# ============= LECTURE 3 EXAMPLE PROBLEM =============
doc.add_heading('Section 2: Example from Lecture 3 Slides', level=1)

para = doc.add_paragraph()
run = para.add_run('Natural Convection of Boundary Layer Flow Along a Vertical Plate')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_heading('Problem Setup (from slides):', level=2)

doc.add_paragraph('Dimensional governing equations (2D):')
doc.add_paragraph()

para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ + ∂v̄/∂ȳ = 0').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂ū/∂x̄) + v̄(∂ū/∂ȳ) = ∂²ū/∂ȳ² + gβ(T - T∞)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = (k/ρCₚ)(∂²T/∂ȳ²)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Boundary Conditions (from slides):', level=2)
doc.add_paragraph('At ȳ = 0: ū = v̄ = 0, T = T∞ + (Tᵥᵥ - T∞)(1 + A sin(πx̄/l))')
doc.add_paragraph('As ȳ → ∞: ū → 0, v̄ → 0, T → T∞')

doc.add_paragraph()
doc.add_heading('Non-dimensional variables (from slides):', level=2)

para = doc.add_paragraph()
para.add_run('x = x̄/l').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('y = Gr^(1/4) · ȳ/l').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('u = (l/ν) · Gr^(-1/2) · ū').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('v = (l/ν) · Gr^(-1/4) · v̄').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('θ = (T - T∞)/(Tᵥᵥ - T∞)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('where Gr = gβ(Tᵥᵥ - T∞)l³/ν²').font.name = 'Courier New'

doc.add_page_break()

# ============= PAGE 2-3 WITH LECTURE INTEGRATION =============
doc.add_heading('Section 3: Non-Dimensionalization Process', level=1)

para = doc.add_paragraph()
run = para.add_run('📖 Your Handnotes + Lecture Slide Methodology')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

if os.path.exists("handnote_page_2.png"):
    doc.add_picture("handnote_page_2.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('From Lecture Slides - Non-dimensional Equations:', level=2)

para = doc.add_paragraph()
para.add_run('∂u/∂x + ∂v/∂y = 0').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Your Handnotes Show the Derivation:', level=2)
doc.add_paragraph('Chain rule applications for each term (see handnote page 2)')

if os.path.exists("handnote_page_3.png"):
    doc.add_picture("handnote_page_3.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============= PAGE 4 WITH LECTURE INTEGRATION =============
doc.add_heading('Section 4: Similarity Transformation', level=1)

para = doc.add_paragraph()
run = para.add_run('📖 From Lecture Slides + Your Derivation')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

if os.path.exists("handnote_page_4.png"):
    doc.add_picture("handnote_page_4.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Lecture Slide Content - Parabolic Transformation:', level=2)

para = doc.add_paragraph()
para.add_run('ψ = x^(3/4) · f(x,η)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('η = x^(-1/4) · y').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('θ = θ(x,η)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Transformed Equations (from slides):', level=2)

para = doc.add_paragraph()
para.add_run("f''' + (3/4)f·f'' - (1/2)(f')² + θ = x(f'·∂f'/∂x - f''·∂f/∂x)").font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run("(1/Pr)θ'' + (3/4)f·θ' = x(f'·∂θ/∂x - θ'·∂f/∂x)").font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('New Boundary Conditions (from slides):', level=2)
doc.add_paragraph('f = f\' = 0, θ = 1 + α sin(πx) at η = 0')
doc.add_paragraph('f\' → 0, θ → 0 as η → ∞')

doc.add_paragraph()
doc.add_heading('From Handnotes:', level=2)
doc.add_paragraph('A (max) = 0.3 (amplitude of sinusoidal variation)')
doc.add_paragraph('P (max) = 0.3 noted in handwriting')

doc.add_page_break()

# ============= LECTURE 3 RESULTS =============
doc.add_heading('Section 5: Code Validation & Results (from Lecture 3)', level=1)

doc.add_heading('From Lecture Slides:', level=2)
doc.add_paragraph('Figure: (a) Shear stress (b) Heat transfer while Pr = 0.7 and α = 0.1')

doc.add_paragraph()
doc.add_heading('Results - Isotherms (from slides):', level=2)
doc.add_paragraph('Figure: Isotherms for (a) α = 0.1 (b) α = 0.2 (c) α = 0.3 while Pr = 0.7')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('⚠ Note: ')
run.bold = True
run.font.color.rgb = RGBColor(255, 100, 0)
para.add_run('The lecture slides show the RESULTS of solving the ODEs numerically. Your handnotes show HOW to derive those ODEs!')

doc.add_page_break()

# ============= INTEGRATION SUMMARY =============
doc.add_heading('Summary: How Handnotes Connect to Lecture 3', level=1)

doc.add_heading('Lecture Flow:', level=2)

doc.add_paragraph('1. Lecture 3 Slides: Introduce boundary layer concept and final equations')
doc.add_paragraph('2. Your Handnotes: Show DETAILED DERIVATION of those equations')
doc.add_paragraph('3. Lecture 3 Slides: Show results (graphs, validation)')

doc.add_paragraph()
doc.add_heading('What Your Professor Expects:', level=2)

para = doc.add_paragraph()
run = para.add_run('✓ Understand the concept (from slides)')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

para = doc.add_paragraph()
run = para.add_run('✓ Derive the equations (from handnotes)')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

para = doc.add_paragraph()
run = para.add_run('✓ Interpret the results (from slides)')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
doc.add_heading('For the Quiz:', level=2)

doc.add_paragraph('Most likely question types:', style='List Bullet')
doc.add_paragraph('1. "Derive the non-dimensional form of..." (handnotes pages 2-3)', style='List Bullet 2')
doc.add_paragraph('2. "What is the boundary layer approximation?" (lecture + handnote page 1)', style='List Bullet 2')
doc.add_paragraph('3. "Define Grashof number and explain its meaning" (lecture slides)', style='List Bullet 2')
doc.add_paragraph('4. "Transform to similarity variables" (handnote page 4)', style='List Bullet 2')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('🎯 Focus: Your handnotes show the PROCESS. Master that process!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(200, 0, 0)

# Save
doc.save('AMCS506_Integrated_Handnotes_and_Lecture3.docx')
print("✓ Created: AMCS506_Integrated_Handnotes_and_Lecture3.docx")
print("\nIntegrated handnotes with Lecture 3 slide content!")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DWT 2012 Problem - Complete Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.add_run("Natural Convection from Vertical Plate").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)

# Problem Statement
doc.add_heading('Problem Statement', level=1)
doc.add_paragraph('Natural convection flow from an isothermal vertical plate with uniform heat source.')

# Dimensional Equations
doc.add_heading('STEP 1: Dimensional Governing Equations', level=1)
para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ + ∂v̄/∂ȳ = 0').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂ū/∂x̄) + v̄(∂ū/∂ȳ) = ν(∂²ū/∂ȳ²) + gβ(T - T∞)').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²) + Q(T - T∞)').font.name = 'Courier New'

doc.add_page_break()

# Dimensionless form
doc.add_heading('STEP 2: Dimensionless Equations', level=1)
para = doc.add_paragraph()
para.add_run('∂u/∂x + ∂v/∂y = 0').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²) + λθ').font.name = 'Courier New'

doc.add_page_break()

# Similarity transformation
doc.add_heading('STEP 3: Similarity Transformation', level=1)
para = doc.add_paragraph()
para.add_run('ψ = x^(3/4) · f(x,η)').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('η = x^(-1/4) · y').font.name = 'Courier New'
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('θ = θ(x,η)').font.name = 'Courier New'

doc.add_page_break()

# Engineering quantities with proper formatting
doc.add_heading('STEP 4: Engineering Quantities', level=1)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Skin Friction Coefficient:')
run.bold = True
run.font.size = Pt(14)

# Create centered paragraph for equation
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Build equation with subscripts/superscripts
run = para.add_run('½ Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('x')
run.font.subscript = True
run.font.size = Pt(11)

run = para.add_run('¹⁄₄')
run.font.superscript = True  
run.font.size = Pt(11)

run = para.add_run(' C')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('f')
run.font.subscript = True
run.font.size = Pt(11)

run = para.add_run(' = (∂u/∂Y)')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('|')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('Y=0')
run.font.subscript = True
run.font.size = Pt(11)

run = para.add_run(' = f\'\'(x,0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

# Nusselt number
para = doc.add_paragraph()
run = para.add_run('Nusselt Number:')
run.bold = True
run.font.size = Pt(14)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = para.add_run('Nu / Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('x')
run.font.subscript = True
run.font.size = Pt(11)

run = para.add_run('¹⁄₄')
run.font.superscript = True
run.font.size = Pt(11)

run = para.add_run(' = -(∂θ/∂Y)')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('|')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

run = para.add_run('Y=0')
run.font.subscript = True
run.font.size = Pt(11)

run = para.add_run(' = -θ\'(x,0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

# Add note
para = doc.add_paragraph()
run = para.add_run('✓ This matches equation (26) in DWT 2012 paper!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)
run.font.size = Pt(12)

# Save
doc.save('DWT_2012_Complete_Solution.docx')
print("✓ Created: DWT_2012_Complete_Solution.docx with proper formatting")

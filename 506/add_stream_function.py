from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document('AMCS506_Complete_Quiz_Preparation.docx')

# Add new section for stream function
doc.add_page_break()
doc.add_heading('DETAILED SECTION: Stream Function & Similarity Transformation', level=1)

para = doc.add_paragraph()
run = para.add_run('⚠ COMPLETE STEP-BY-STEP SOLUTION')
run.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_heading('Part 1: What is a Stream Function?', level=2)

doc.add_heading('The Problem:', level=3)
doc.add_paragraph('We have 2 equations (continuity + momentum) with 2 unknowns (u, v).')
doc.add_paragraph('Stream function is a mathematical trick to reduce this to 1 equation with 1 unknown!')

doc.add_paragraph()
doc.add_heading('Definition:', level=3)

para = doc.add_paragraph()
run = para.add_run('u = ∂ψ/∂y')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('v = -∂ψ/∂x')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
doc.add_heading('Why This Works (Proof):', level=3)

para = doc.add_paragraph()
para.add_run('Continuity equation: ∂u/∂x + ∂v/∂y = 0').bold = True

doc.add_paragraph()
doc.add_paragraph('Substitute u and v:')
para = doc.add_paragraph()
para.add_run('∂u/∂x = ∂/∂x(∂ψ/∂y) = ∂²ψ/∂x∂y').font.name = 'Courier New'

para = doc.add_paragraph()
para.add_run('∂v/∂y = ∂/∂y(-∂ψ/∂x) = -∂²ψ/∂y∂x').font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Add them:')
para = doc.add_paragraph()
para.add_run('∂²ψ/∂x∂y + (-∂²ψ/∂y∂x) = ∂²ψ/∂x∂y - ∂²ψ/∂x∂y = 0 ✓').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Continuity is AUTOMATICALLY satisfied!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

doc.add_heading('Part 2: Similarity Transformation', level=2)

doc.add_heading('The Goal:', level=3)
doc.add_paragraph('Convert PDEs (partial differential equations) to ODEs (ordinary differential equations)')
doc.add_paragraph('PDEs depend on both x and y → hard to solve')
doc.add_paragraph('ODEs depend on one variable η → easier to solve!')

doc.add_paragraph()
doc.add_heading('Step 1: Define Similarity Variables', level=3)

para = doc.add_paragraph()
run = para.add_run('ψ = x^(3/4) · f(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('η = x^(-1/4) · y')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('θ = θ(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
doc.add_heading('Why These Specific Powers?', level=3)
doc.add_paragraph('From similarity analysis of natural convection boundary layers:')
doc.add_paragraph('x^(3/4): Comes from balance of buoyancy and viscous forces', style='List Bullet')
doc.add_paragraph('x^(-1/4): Makes η dimensionless and O(1) in boundary layer', style='List Bullet')
doc.add_paragraph('These are NOT arbitrary - they come from scaling analysis!', style='List Bullet')

doc.add_page_break()

doc.add_heading('Part 3: Calculate Velocity Components', level=2)

doc.add_heading('Step 2: Find u = ∂ψ/∂y', level=3)

para = doc.add_paragraph()
para.add_run('Given: ψ = x^(3/4) · f(x,η)').bold = True

doc.add_paragraph()
doc.add_paragraph('Use chain rule:')
para = doc.add_paragraph()
para.add_run('u = ∂ψ/∂y = ∂ψ/∂η · ∂η/∂y').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Calculate ∂η/∂y:').bold = True
para = doc.add_paragraph()
para.add_run('η = x^(-1/4) · y').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('∂η/∂y = x^(-1/4)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Calculate ∂ψ/∂η:').bold = True
para = doc.add_paragraph()
para.add_run('ψ = x^(3/4) · f(x,η)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('∂ψ/∂η = x^(3/4) · ∂f/∂η = x^(3/4) · f\'').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Combine:').bold = True
para = doc.add_paragraph()
para.add_run('u = x^(3/4) · f\' · x^(-1/4)').font.name = 'Courier New'
para = doc.add_paragraph()
run = para.add_run('u = x^(1/2) · f\'(x,η)')
run.font.name = 'Courier New'
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
doc.add_heading('Step 3: Find v = -∂ψ/∂x', level=3)

para = doc.add_paragraph()
para.add_run('This is more complex! Need product rule AND chain rule.').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('v = -∂ψ/∂x = -∂/∂x[x^(3/4) · f(x,η)]').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Product rule:').bold = True
para = doc.add_paragraph()
para.add_run('= -[∂(x^(3/4))/∂x · f + x^(3/4) · ∂f/∂x]').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('First term:').bold = True
para = doc.add_paragraph()
para.add_run('∂(x^(3/4))/∂x = (3/4)x^(-1/4)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Second term (chain rule):').bold = True
para = doc.add_paragraph()
para.add_run('∂f/∂x = ∂f/∂x|_η + ∂f/∂η · ∂η/∂x').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Calculate ∂η/∂x:').bold = True
para = doc.add_paragraph()
para.add_run('η = x^(-1/4) · y').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('∂η/∂x = (-1/4)x^(-5/4) · y = (-1/4)x^(-5/4) · x^(1/4)η').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('     = (-1/4)x^(-1) · η').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Combine everything:').bold = True
para = doc.add_paragraph()
para.add_run('v = -[(3/4)x^(-1/4)·f + x^(3/4)·(∂f/∂x|_η + f\'·(-1/4)x^(-1)·η)]').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Simplify:').bold = True
para = doc.add_paragraph()
run = para.add_run('v = -(3/4)x^(-1/4)·f - x^(3/4)·∂f/∂x + (1/4)x^(-1/4)·η·f\'')
run.font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('v = (1/4)x^(-1/4)·[ηf\' - 3f] - x^(3/4)·∂f/∂x')
run.font.name = 'Courier New'
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_page_break()

doc.add_heading('Part 4: Calculate Derivatives for Momentum Equation', level=2)

doc.add_heading('Step 4: Calculate ∂u/∂x', level=3)

para = doc.add_paragraph()
para.add_run('u = x^(1/2) · f\'(x,η)').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂u/∂x = ∂/∂x[x^(1/2) · f\']').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Product rule:').bold = True
para = doc.add_paragraph()
para.add_run('= (1/2)x^(-1/2)·f\' + x^(1/2)·∂f\'/∂x').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Chain rule for ∂f\'/∂x:').bold = True
para = doc.add_paragraph()
para.add_run('∂f\'/∂x = ∂²f/∂x²|_η + ∂²f/∂η∂x · ∂η/∂x').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('        = ∂²f/∂x²|_η + f\'\'·(-1/4)x^(-1)·η').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('∂u/∂x = (1/2)x^(-1/2)·f\' + x^(1/2)·∂²f/∂x² - (1/4)x^(-1/2)·η·f\'\'')
run.font.name = 'Courier New'
run.bold = True

doc.add_paragraph()
doc.add_heading('Step 5: Calculate ∂u/∂y', level=3)

para = doc.add_paragraph()
para.add_run('u = x^(1/2) · f\'').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂u/∂y = x^(1/2) · ∂f\'/∂y').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Chain rule:').bold = True
para = doc.add_paragraph()
para.add_run('∂f\'/∂y = ∂f\'/∂η · ∂η/∂y = f\'\' · x^(-1/4)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('∂u/∂y = x^(1/2) · f\'\' · x^(-1/4) = x^(1/4) · f\'\'')
run.font.name = 'Courier New'
run.bold = True

doc.add_paragraph()
doc.add_heading('Step 6: Calculate ∂²u/∂y²', level=3)

para = doc.add_paragraph()
para.add_run('∂u/∂y = x^(1/4) · f\'\'').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂²u/∂y² = ∂/∂y[x^(1/4) · f\'\']').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= x^(1/4) · ∂f\'\'/∂y').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= x^(1/4) · f\'\'\' · ∂η/∂y').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('∂²u/∂y² = x^(1/4) · f\'\'\' · x^(-1/4) = f\'\'\'')
run.font.name = 'Courier New'
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Notice: The x terms cancel! This is key to similarity!')
run.italic = True
run.font.color.rgb = RGBColor(200, 0, 0)

doc.add_page_break()

doc.add_heading('Part 5: Substitute into Momentum Equation', level=2)

doc.add_heading('Step 7: Original Dimensionless Momentum Equation', level=3)

para = doc.add_paragraph()
run = para.add_run('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
doc.add_heading('Step 8: Substitute All Derivatives', level=3)

para = doc.add_paragraph()
para.add_run('Left side - Term 1: u(∂u/∂x)').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂u/∂x) = x^(1/2)·f\' · [(1/2)x^(-1/2)·f\' + x^(1/2)·∂²f/∂x² - (1/4)x^(-1/2)·η·f\'\']').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (1/2)f\'² + x·f\'·∂²f/∂x² - (1/4)η·f\'·f\'\'').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Left side - Term 2: v(∂u/∂y)').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('v(∂u/∂y) = [(1/4)x^(-1/4)·(ηf\' - 3f) - x^(3/4)·∂f/∂x] · x^(1/4)·f\'\'').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (1/4)(ηf\' - 3f)·f\'\' - x·f\'\'·∂f/∂x').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (1/4)η·f\'·f\'\' - (3/4)f·f\'\' - x·f\'\'·∂f/∂x').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 9: Combine Left Side', level=3)

para = doc.add_paragraph()
para.add_run('u(∂u/∂x) + v(∂u/∂y) =').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (1/2)f\'² + x·f\'·∂²f/∂x² - (1/4)η·f\'·f\'\' + (1/4)η·f\'·f\'\' - (3/4)f·f\'\' - x·f\'\'·∂f/∂x').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Notice: -(1/4)η·f\'·f\'\' and +(1/4)η·f\'·f\'\' cancel!').italic = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (1/2)f\'² - (3/4)f·f\'\' + x[f\'·∂²f/∂x² - f\'\'·∂f/∂x]').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 10: Final ODE Form', level=3)

para = doc.add_paragraph()
para.add_run('Set equal to right side:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('(1/2)f\'² - (3/4)f·f\'\' + x[f\'·∂²f/∂x² - f\'\'·∂f/∂x] = f\'\'\' + θ').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Rearrange:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('f\'\'\' + (3/4)f·f\'\' - (1/2)f\'² + θ = x[f\'·∂²f/∂x² - f\'\'·∂f/∂x]')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('✓ THIS IS THE FINAL ODE!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

doc.add_heading('Part 6: Summary & Key Points', level=2)

doc.add_heading('What We Accomplished:', level=3)

doc.add_paragraph('1. Reduced 2 equations (continuity + momentum) to 1 equation', style='List Number')
doc.add_paragraph('2. Converted PDEs (u, v functions of x, y) to ODE (f function of η)', style='List Number')
doc.add_paragraph('3. Made problem easier to solve numerically', style='List Number')

doc.add_paragraph()
doc.add_heading('Key Formulas to Remember:', level=3)

para = doc.add_paragraph()
para.add_run('Stream function:').bold = True
para = doc.add_paragraph()
para.add_run('ψ = x^(3/4) · f(x,η),  η = x^(-1/4) · y').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Velocities:').bold = True
para = doc.add_paragraph()
para.add_run('u = x^(1/2) · f\'').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('v = (1/4)x^(-1/4) · [ηf\' - 3f] - x^(3/4) · ∂f/∂x').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Final ODE:').bold = True
para = doc.add_paragraph()
para.add_run('f\'\'\' + (3/4)f·f\'\' - (1/2)f\'² + θ = x[f\'·∂²f/∂x² - f\'\'·∂f/∂x]').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Practice Problems:', level=3)

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 1: Verify u = ∂ψ/∂y gives u = x^(1/2)·f\'').bold = True
doc.add_paragraph('Hint: Use chain rule with η = x^(-1/4)·y')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 2: Calculate ∂²u/∂y² step-by-step').bold = True
doc.add_paragraph('Answer: f\'\'\' (x terms cancel!)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 3: Why does similarity work?').bold = True
doc.add_paragraph('Answer: Because x-dependence separates out, leaving ODE in η')

doc.add_paragraph()
doc.add_heading('Why This Matters for Quiz:', level=3)

doc.add_paragraph('Your professor showed this derivation in handnotes pages 4-7')
doc.add_paragraph('Understanding the process is more important than memorizing the final form')
doc.add_paragraph('Focus on: chain rule, product rule, and how x terms cancel')

# Save
doc.save('AMCS506_Complete_Quiz_Preparation.docx')
print("✓ Updated: AMCS506_Complete_Quiz_Preparation.docx")
print("\nAdded complete stream function derivation!")

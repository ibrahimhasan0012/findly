from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document('AMCS506_Complete_Quiz_Preparation.docx')

# Add new section for energy equation
doc.add_page_break()
doc.add_heading('DETAILED SECTION: Energy Equation Derivation & Non-Dimensionalization', level=1)

para = doc.add_paragraph()
run = para.add_run('⚠ ADDED: Complete step-by-step energy equation treatment')
run.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph()
doc.add_heading('Part 1: Energy Conservation - First Principles', level=2)

doc.add_heading('Physical Principle:', level=3)
para = doc.add_paragraph()
para.add_run('First Law of Thermodynamics: ').bold = True
para.add_run('Rate of change of energy = Heat added + Work done')

doc.add_paragraph()
doc.add_heading('For a Fluid Element:', level=3)
doc.add_paragraph('Energy can change due to:', style='List Bullet')
doc.add_paragraph('1. Convection (fluid carrying energy)', style='List Bullet 2')
doc.add_paragraph('2. Conduction (heat diffusion)', style='List Bullet 2')
doc.add_paragraph('3. Internal heat generation/absorption', style='List Bullet 2')

doc.add_paragraph()
doc.add_heading('Dimensional Energy Equation:', level=2)

para = doc.add_paragraph()
run = para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²) + Q/(ρCₚ)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
doc.add_heading('Term-by-Term Breakdown:', level=3)

para = doc.add_paragraph()
para.add_run('LEFT SIDE - Convective Heat Transfer:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄): ').bold = True
para.add_run('Heat carried by fluid moving in x-direction')
doc.add_paragraph('Physical meaning: Hot fluid moving downstream carries thermal energy', style='List Bullet 2')
doc.add_paragraph('Example: Hot air from a heater flowing across a room', style='List Bullet 2')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('v̄(∂T/∂ȳ): ').bold = True
para.add_run('Heat carried by fluid moving in y-direction')
doc.add_paragraph('Physical meaning: Hot fluid rising (or sinking) carries energy', style='List Bullet 2')
doc.add_paragraph('Example: Hot air rising from a candle flame', style='List Bullet 2')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('RIGHT SIDE - Diffusive Heat Transfer:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('α(∂²T/∂ȳ²): ').bold = True
para.add_run('Heat conduction (Fourier\'s law)')
doc.add_paragraph('α = k/(ρCₚ) = thermal diffusivity [m²/s]', style='List Bullet 2')
doc.add_paragraph('Physical meaning: Heat spreading by molecular collisions', style='List Bullet 2')
doc.add_paragraph('Example: Heat spreading through a metal rod', style='List Bullet 2')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Q/(ρCₚ): ').bold = True
para.add_run('Internal heat generation/absorption')
doc.add_paragraph('Q > 0: Heat generation (exothermic reaction, electrical heating)', style='List Bullet 2')
doc.add_paragraph('Q < 0: Heat absorption (endothermic reaction, cooling)', style='List Bullet 2')
doc.add_paragraph('Q = 0: No internal heat source (most common case)', style='List Bullet 2')

doc.add_page_break()

doc.add_heading('Part 2: Non-Dimensionalization of Energy Equation', level=2)

para = doc.add_paragraph()
run = para.add_run('🎯 STEP-BY-STEP DERIVATION')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
doc.add_heading('Step 1: Define Dimensionless Variables', level=3)

doc.add_paragraph('Coordinates:')
para = doc.add_paragraph()
para.add_run('x = x̄/l,  y = (Gr^(1/4)/l)·ȳ').font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Velocities:')
para = doc.add_paragraph()
para.add_run('u = (l/ν)·Gr^(-1/2)·ū,  v = (l/ν)·Gr^(-1/4)·v̄').font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Temperature:')
para = doc.add_paragraph()
para.add_run('θ = (T - T∞)/(Tᵥᵥ - T∞) = (T - T∞)/ΔT').font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Therefore:')
para = doc.add_paragraph()
para.add_run('T = T∞ + ΔT·θ').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 2: Transform ∂T/∂x̄ (First Term)', level=3)

para = doc.add_paragraph()
para.add_run('Using chain rule:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂T/∂x̄ = ∂(T∞ + ΔT·θ)/∂x̄').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Since T∞ is constant:').italic = True
para = doc.add_paragraph()
para.add_run('∂T/∂x̄ = ΔT·(∂θ/∂x̄)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Apply chain rule:').bold = True
para = doc.add_paragraph()
para.add_run('∂θ/∂x̄ = (∂θ/∂x)·(∂x/∂x̄)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Since x = x̄/l:').italic = True
para = doc.add_paragraph()
para.add_run('∂x/∂x̄ = 1/l').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Therefore:').bold = True
para = doc.add_paragraph()
para.add_run('∂T/∂x̄ = ΔT·(∂θ/∂x)·(1/l) = (ΔT/l)·(∂θ/∂x)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 3: Transform ∂T/∂ȳ (Second Term)', level=3)

para = doc.add_paragraph()
para.add_run('Using chain rule:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂T/∂ȳ = ΔT·(∂θ/∂ȳ)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂θ/∂ȳ = (∂θ/∂y)·(∂y/∂ȳ)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Since y = (Gr^(1/4)/l)·ȳ:').italic = True
para = doc.add_paragraph()
para.add_run('∂y/∂ȳ = Gr^(1/4)/l').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Therefore:').bold = True
para = doc.add_paragraph()
para.add_run('∂T/∂ȳ = ΔT·(∂θ/∂y)·(Gr^(1/4)/l) = (ΔT·Gr^(1/4)/l)·(∂θ/∂y)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 4: Transform ∂²T/∂ȳ² (Diffusion Term)', level=3)

para = doc.add_paragraph()
para.add_run('Take derivative of ∂T/∂ȳ with respect to ȳ:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('∂²T/∂ȳ² = ∂/∂ȳ[(ΔT·Gr^(1/4)/l)·(∂θ/∂y)]').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (ΔT·Gr^(1/4)/l)·∂/∂ȳ(∂θ/∂y)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (ΔT·Gr^(1/4)/l)·(∂²θ/∂y²)·(∂y/∂ȳ)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (ΔT·Gr^(1/4)/l)·(∂²θ/∂y²)·(Gr^(1/4)/l)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('= (ΔT·Gr^(1/2)/l²)·(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_page_break()

doc.add_heading('Step 5: Substitute into Energy Equation', level=3)

para = doc.add_paragraph()
para.add_run('Original equation:').bold = True
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Substitute dimensional velocities:').bold = True
para = doc.add_paragraph()
para.add_run('ū = (ν/l)·Gr^(1/2)·u').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('v̄ = (ν/l)·Gr^(1/4)·v').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('LEFT SIDE:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) = (ν/l)·Gr^(1/2)·u · (ΔT/l)·(∂θ/∂x)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('           = (ν·ΔT/l²)·Gr^(1/2)·u·(∂θ/∂x)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('v̄(∂T/∂ȳ) = (ν/l)·Gr^(1/4)·v · (ΔT·Gr^(1/4)/l)·(∂θ/∂y)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('           = (ν·ΔT/l²)·Gr^(1/2)·v·(∂θ/∂y)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Combined LEFT SIDE:').bold = True
para = doc.add_paragraph()
para.add_run('(ν·ΔT/l²)·Gr^(1/2)·[u(∂θ/∂x) + v(∂θ/∂y)]').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('RIGHT SIDE:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('α(∂²T/∂ȳ²) = α·(ΔT·Gr^(1/2)/l²)·(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
doc.add_heading('Step 6: Simplify and Cancel', level=3)

para = doc.add_paragraph()
para.add_run('Set LEFT = RIGHT:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('(ν·ΔT/l²)·Gr^(1/2)·[u(∂θ/∂x) + v(∂θ/∂y)] = α·(ΔT·Gr^(1/2)/l²)·(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Divide both sides by (ΔT·Gr^(1/2)/l²):').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('ν·[u(∂θ/∂x) + v(∂θ/∂y)] = α·(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Divide both sides by ν:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (α/ν)·(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Since Pr = ν/α, we have α/ν = 1/Pr:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)·(∂²θ/∂y²)')
run.font.name = 'Courier New'
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('✓ FINAL DIMENSIONLESS ENERGY EQUATION!')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

doc.add_heading('Part 3: Physical Interpretation', level=2)

doc.add_heading('What Does 1/Pr Mean?', level=3)

doc.add_paragraph('Pr = ν/α = (momentum diffusivity)/(thermal diffusivity)')
doc.add_paragraph('1/Pr = α/ν = (thermal diffusivity)/(momentum diffusivity)')

doc.add_paragraph()
doc.add_paragraph('For air (Pr = 0.7):')
doc.add_paragraph('1/Pr = 1/0.7 ≈ 1.43', style='List Bullet')
doc.add_paragraph('Heat diffuses 1.43× faster than momentum', style='List Bullet')
doc.add_paragraph('Thermal boundary layer is thicker', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('For water (Pr = 7):')
doc.add_paragraph('1/Pr = 1/7 ≈ 0.14', style='List Bullet')
doc.add_paragraph('Heat diffuses 7× slower than momentum', style='List Bullet')
doc.add_paragraph('Thermal boundary layer is thinner', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Energy Equation with Heat Source:', level=3)

para = doc.add_paragraph()
para.add_run('If Q ≠ 0:').bold = True
para = doc.add_paragraph()
para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)·(∂²θ/∂y²) + λθ').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('where λ = QL²/ν is the heat generation parameter').italic = True

doc.add_paragraph()
doc.add_heading('Practice Problems:', level=2)

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 1: Derive ∂T/∂ȳ step-by-step').bold = True
doc.add_paragraph('Start with T = T∞ + ΔT·θ and y = (Gr^(1/4)/l)·ȳ')
doc.add_paragraph('Answer: ∂T/∂ȳ = (ΔT·Gr^(1/4)/l)·(∂θ/∂y)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 2: If Pr = 0.7, what is 1/Pr?').bold = True
doc.add_paragraph('Answer: 1/Pr = 1.43 (heat diffuses faster)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Problem 3: Write energy equation with heat generation').bold = True
doc.add_paragraph('Answer: u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²) + λθ')

doc.add_paragraph()
doc.add_heading('Summary - Energy Equation Transformation:', level=2)

para = doc.add_paragraph()
run = para.add_run('DIMENSIONAL → DIMENSIONLESS')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Before:').bold = True
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('After:').bold = True
para = doc.add_paragraph()
para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Key: All ΔT, Gr, l, ν factors cancel out, leaving only Pr!')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

# Save updated document
doc.save('AMCS506_Complete_Quiz_Preparation.docx')
print("✓ Updated: AMCS506_Complete_Quiz_Preparation.docx")
print("\nAdded detailed energy equation derivation section!")

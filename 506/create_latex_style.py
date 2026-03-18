from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_equation(paragraph, latex_text):
    """Insert a LaTeX-style equation using OMML (Office Math Markup Language)"""
    # This creates a proper Word equation field
    run = paragraph.add_run()
    
    # Create equation element
    math = OxmlElement('m:oMath')
    math_para = OxmlElement('m:oMathPara')
    
    # For now, we'll use a formatted text approach with Cambria Math
    # Word equations are complex, so we'll create visually similar output
    return run

doc = Document()

# Title
title = doc.add_heading('DWT 2012 Problem - Complete Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Cambria'

doc.add_paragraph()

# Momentum equation - LaTeX style
doc.add_heading('Momentum (boundary layer approximation):', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Build equation piece by piece with Cambria Math
run = para.add_run('ū')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂ū')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('/')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('x̄')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(') + ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('v̄')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂ū')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('/')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('ȳ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(') = ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('ν')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂²ū')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('/')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('ȳ²')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(') + ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('g')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('β')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('T')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(' − ')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('T')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run('∞,')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)
run.font.italic = True

run = para.add_run(')')
run.font.name = 'Cambria Math'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Engineering quantities - equation (26) style
doc.add_heading('Engineering Quantities (Equation 26):', level=1)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'Cambria'

doc.add_paragraph()

# Skin friction
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = para.add_run('½')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run(' Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run('1/4')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.superscript = True

run = para.add_run(' ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('C')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('f')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run(' = ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('u')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('/')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(')')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('|')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('Y=0')
run.font.name = 'Cambria Math'
run.font.size = Pt(9)
run.font.subscript = True

run = para.add_run(' = ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('f')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('′′')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(', 0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

doc.add_paragraph()

# Nusselt number
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = para.add_run('Nu')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(' / ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Gr')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.subscript = True

run = para.add_run('1/4')
run.font.name = 'Cambria Math'
run.font.size = Pt(10)
run.font.superscript = True

run = para.add_run(' = −')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('θ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('/')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('∂')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('Y')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(')')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('|')
run.font.name = 'Cambria Math'
run.font.size = Pt(16)

run = para.add_run('Y=0')
run.font.name = 'Cambria Math'
run.font.size = Pt(9)
run.font.subscript = True

run = para.add_run(' = −')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('θ')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run('′')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('(')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

run = para.add_run('x')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)
run.font.italic = True

run = para.add_run(', 0)')
run.font.name = 'Cambria Math'
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('✓ LaTeX-style formatting with Cambria Math font')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)
run.font.name = 'Cambria'

# Save
doc.save('DWT_2012_LaTeX_Style.docx')
print("✓ Created: DWT_2012_LaTeX_Style.docx")
print("\nLaTeX-style equations with Cambria Math font!")

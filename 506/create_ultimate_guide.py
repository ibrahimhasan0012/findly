from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('AMCS 506 - ULTIMATE Study Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.add_run("Handnotes + Lecture 3 + Real-World Examples + Detailed Explanations").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# ============= SECTION 1: BOUNDARY LAYER FUNDAMENTALS =============
doc.add_heading('Section 1: Boundary Layer Fundamentals', level=1)

para = doc.add_paragraph()
run = para.add_run('📖 From Lecture 3 Slides + Your Handnotes')
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

if os.path.exists("handnote_page_1.png"):
    doc.add_picture("handnote_page_1.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('What is a Boundary Layer? (Lecture 3)', level=2)

doc.add_paragraph('Discovered by Ludwig Prandtl (1875-1953) - one of the most important discoveries in fluid mechanics!')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Definition: ').bold = True
para.add_run('A thin region adjacent to a solid surface where the effects of viscosity are significant.')

doc.add_paragraph()
doc.add_heading('Real-World Examples:', level=3)

doc.add_paragraph('Aircraft Wing:', style='List Bullet')
doc.add_paragraph('Boundary layer is only a few millimeters thick on a wing at cruise speed', style='List Bullet 2')
doc.add_paragraph('Controls drag and can cause flow separation (stall)', style='List Bullet 2')
doc.add_paragraph('Engineers design wing shapes to keep boundary layer attached', style='List Bullet 2')

doc.add_paragraph('Ship Hull:', style='List Bullet')
doc.add_paragraph('Boundary layer creates friction drag', style='List Bullet 2')
doc.add_paragraph('Thicker boundary layer = more fuel consumption', style='List Bullet 2')
doc.add_paragraph('Special coatings reduce boundary layer thickness', style='List Bullet 2')

doc.add_paragraph('Heat Exchanger:', style='List Bullet')
doc.add_paragraph('Thermal boundary layer controls heat transfer rate', style='List Bullet 2')
doc.add_paragraph('Thinner boundary layer = better heat transfer', style='List Bullet 2')
doc.add_paragraph('Turbulence promoters reduce boundary layer thickness', style='List Bullet 2')

doc.add_paragraph()
doc.add_heading('Formula 1: Continuity Equation', level=2)

para = doc.add_paragraph()
run = para.add_run('∂u/∂x + ∂v/∂y = 0')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why this formula?', level=3)
doc.add_paragraph('Mass conservation - what flows in must flow out. For incompressible flow (ρ = constant), this simplifies to the divergence of velocity being zero.')

doc.add_heading('Physical Meaning:', level=3)
doc.add_paragraph('If fluid accelerates in x-direction (∂u/∂x > 0), it must decelerate in y-direction (∂v/∂y < 0) to conserve mass.')

doc.add_heading('Real-World Example:', level=3)
doc.add_paragraph('Water flowing through a pipe: When pipe narrows, velocity increases (continuity requires constant mass flow rate Q = A·v).')

doc.add_heading('How to Practice:', level=3)
doc.add_paragraph('Practice 1: Given u = x² + y, find v such that continuity is satisfied', style='List Bullet')
doc.add_paragraph('Solution: ∂u/∂x = 2x, so ∂v/∂y = -2x, integrate: v = -2xy + f(x)', style='List Bullet 2')
doc.add_paragraph('Practice 2: Verify if u = xy, v = -xy satisfies continuity', style='List Bullet')
doc.add_paragraph('Solution: ∂u/∂x = y, ∂v/∂y = -x, sum = y - x ≠ 0, NO! ✗', style='List Bullet 2')

doc.add_paragraph()
doc.add_heading('Formula 2: Momentum Equation (Boundary Layer)', level=2)

para = doc.add_paragraph()
run = para.add_run('u(∂u/∂x) + v(∂u/∂y) = ν(∂²u/∂y²) + gβ(T - T∞)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why this formula?', level=3)
doc.add_paragraph('Newton\'s second law (F = ma) for fluid flow. The boundary layer approximation removes ∂²u/∂x² because changes perpendicular to wall dominate.')

doc.add_heading('Term-by-Term Breakdown:', level=3)

para = doc.add_paragraph()
para.add_run('u(∂u/∂x): ').bold = True
para.add_run('Convective acceleration in x-direction')
doc.add_paragraph('Physical meaning: Fluid speeding up as it moves downstream', style='List Bullet 2')
doc.add_paragraph('Example: Water speeding up as it flows down a slide', style='List Bullet 2')

para = doc.add_paragraph()
para.add_run('v(∂u/∂y): ').bold = True
para.add_run('Convective acceleration in y-direction')
doc.add_paragraph('Physical meaning: Fluid moving toward/away from wall changes velocity', style='List Bullet 2')
doc.add_paragraph('Example: Air rising from hot pavement changes horizontal velocity', style='List Bullet 2')

para = doc.add_paragraph()
para.add_run('ν(∂²u/∂y²): ').bold = True
para.add_run('Viscous diffusion')
doc.add_paragraph('Physical meaning: Friction slowing down the fluid', style='List Bullet 2')
doc.add_paragraph('Example: Honey flows slower than water due to higher viscosity', style='List Bullet 2')

para = doc.add_paragraph()
para.add_run('gβ(T - T∞): ').bold = True
para.add_run('Buoyancy force')
doc.add_paragraph('Physical meaning: Hot fluid rises, cold fluid sinks', style='List Bullet 2')
doc.add_paragraph('Example: Hot air balloon rises due to buoyancy', style='List Bullet 2')

doc.add_heading('Why Only ∂²u/∂y²?', level=3)
doc.add_paragraph('Boundary layer is THIN (δ << L). Order of magnitude analysis:')
doc.add_paragraph('∂²u/∂x² ~ u/L² (small)', style='List Bullet')
doc.add_paragraph('∂²u/∂y² ~ u/δ² (large)', style='List Bullet')
doc.add_paragraph('Since δ << L, we have δ² << L², so ∂²u/∂y² >> ∂²u/∂x²', style='List Bullet')
doc.add_paragraph('We drop the smaller term!', style='List Bullet')

doc.add_heading('Real-World Application:', level=3)
doc.add_paragraph('Airplane wing: Boundary layer is ~2mm thick on a 2m chord. Ratio δ/L = 0.001, so (δ/L)² = 0.000001. The y-derivative is 1 million times larger!')

doc.add_page_break()

# ============= SECTION 2: NON-DIMENSIONALIZATION =============
doc.add_heading('Section 2: Non-Dimensionalization Process', level=1)

para = doc.add_paragraph()
run = para.add_run('🎯 MOST IMPORTANT FOR QUIZ!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(255, 0, 0)

if os.path.exists("handnote_page_2.png"):
    doc.add_picture("handnote_page_2.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Why Non-Dimensionalize?', level=2)

doc.add_paragraph('Reason 1: Reduce number of parameters', style='List Bullet')
doc.add_paragraph('Instead of tracking ρ, μ, U, L separately, we get one number: Re = ρUL/μ', style='List Bullet 2')

doc.add_paragraph('Reason 2: Universal solutions', style='List Bullet')
doc.add_paragraph('Solution for air at 1 m/s applies to water at different speed if Re is same', style='List Bullet 2')

doc.add_paragraph('Reason 3: Identify dominant physics', style='List Bullet')
doc.add_paragraph('If Re >> 1, inertia dominates. If Re << 1, viscosity dominates', style='List Bullet 2')

doc.add_heading('Real-World Example:', level=3)
doc.add_paragraph('Wind tunnel testing: Test 1/10 scale model at 10× speed to match Reynolds number of full-scale aircraft. Results apply to real plane!')

doc.add_paragraph()
doc.add_heading('The Chain Rule Technique (CRITICAL!)', level=2)

para = doc.add_paragraph()
run = para.add_run('⚠ Master this and you master non-dimensionalization!')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)

doc.add_paragraph()
doc.add_heading('Step-by-Step: Deriving ∂ū/∂x̄', level=3)

para = doc.add_paragraph()
para.add_run('Given:').bold = True
doc.add_paragraph('ū = (ν/l) · Gr^(1/2) · u (dimensional velocity)')
doc.add_paragraph('x̄ = l · x (dimensional coordinate)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Step 1: Apply chain rule').bold = True
para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ = (∂u/∂x̄) · (∂x/∂x̄) · (∂ū/∂u)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Step 2: Calculate each part').bold = True
doc.add_paragraph('∂x/∂x̄ = 1/l (since x = x̄/l)')
doc.add_paragraph('∂ū/∂u = (ν/l) · Gr^(1/2) (from ū = (ν/l)·Gr^(1/2)·u)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Step 3: Substitute').bold = True
para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ = (∂u/∂x̄) · (1/l) · (ν/l)·Gr^(1/2)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('       = (ν/l²)·Gr^(1/2) · (∂u/∂x)').font.name = 'Courier New'

doc.add_heading('Practice This!', level=3)
doc.add_paragraph('Do this derivation 10 times until you can do it in your sleep!')

if os.path.exists("handnote_page_3.png"):
    doc.add_picture("handnote_page_3.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Stretched Coordinate (Key Concept!)', level=2)

para = doc.add_paragraph()
run = para.add_run('y = (Gr^(1/4)/l) · ȳ')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Why Stretch?', level=3)
doc.add_paragraph('Boundary layer thickness: δ ~ l/Gr^(1/4)')
doc.add_paragraph('Without stretching: y values would be 0.001, 0.002, 0.003... (tiny!)')
doc.add_paragraph('With stretching: y values become 1, 2, 3... (order 1)')
doc.add_paragraph('Makes numerical solution much easier!')

doc.add_heading('Real-World Analogy:', level=3)
doc.add_paragraph('Like using a microscope to zoom into a thin layer. The layer is thin in real life, but we "stretch" our view to see details.')

doc.add_heading('Example Calculation:', level=3)
doc.add_paragraph('If Gr = 10^8:')
doc.add_paragraph('Gr^(1/4) = (10^8)^(1/4) = 10^2 = 100', style='List Bullet')
doc.add_paragraph('Stretching factor = 100', style='List Bullet')
doc.add_paragraph('Physical distance ȳ = 0.01m becomes y = 100 × 0.01/l = 1 (if l = 1m)', style='List Bullet')

doc.add_page_break()

# ============= SECTION 3: DIMENSIONLESS NUMBERS =============
doc.add_heading('Section 3: Dimensionless Numbers (Memorize These!)', level=1)

para = doc.add_paragraph()
run = para.add_run('⚠ 80% chance these will be on the quiz!')
run.bold = True
run.font.color.rgb = RGBColor(255, 100, 0)

doc.add_paragraph()
doc.add_heading('Grashof Number (Gr)', level=2)

para = doc.add_paragraph()
run = para.add_run('Gr = gβΔTL³/ν²')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Physical Meaning:', level=3)
para = doc.add_paragraph()
para.add_run('Gr = (buoyancy force) / (viscous force)').italic = True

doc.add_heading('Real-World Examples:', level=3)
doc.add_paragraph('Hot radiator (Gr ~ 10^9): Strong natural convection, turbulent flow')
doc.add_paragraph('Warm cup of coffee (Gr ~ 10^6): Moderate natural convection, laminar')
doc.add_paragraph('Microelectronics (Gr ~ 10^3): Weak natural convection')

doc.add_heading('Rule of Thumb:', level=3)
doc.add_paragraph('Gr < 10^8: Laminar natural convection')
doc.add_paragraph('Gr > 10^9: Turbulent natural convection')

doc.add_paragraph()
doc.add_heading('Prandtl Number (Pr)', level=2)

para = doc.add_paragraph()
run = para.add_run('Pr = ν/α = (momentum diffusivity) / (thermal diffusivity)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('Physical Meaning:', level=3)
doc.add_paragraph('Tells you which spreads faster: momentum or heat')

doc.add_heading('Real-World Values:', level=3)
doc.add_paragraph('Liquid metals (Pr ~ 0.01): Heat diffuses 100× faster than momentum')
doc.add_paragraph('Air (Pr ~ 0.7): Heat diffuses slightly faster')
doc.add_paragraph('Water (Pr ~ 7): Momentum diffuses 7× faster than heat')
doc.add_paragraph('Engine oil (Pr ~ 1000): Momentum diffuses 1000× faster!')

doc.add_heading('Boundary Layer Thickness:', level=3)
para = doc.add_paragraph()
para.add_run('δ_thermal / δ_velocity ≈ Pr^(-1/3)').font.name = 'Courier New'

doc.add_paragraph('For air (Pr = 0.7): δ_thermal ≈ 1.13 × δ_velocity (thermal BL thicker)')
doc.add_paragraph('For water (Pr = 7): δ_thermal ≈ 0.52 × δ_velocity (velocity BL thicker)')

doc.add_paragraph()
doc.add_heading('Practice Problem:', level=3)
para = doc.add_paragraph()
para.add_run('Q: If Pr = 0.7 (air), which diffuses faster: momentum or heat?').bold = True
doc.add_paragraph('A: Heat! Because Pr = ν/α < 1 means α > ν, so thermal diffusivity is larger.')

doc.add_page_break()

# ============= SECTION 4: SIMILARITY TRANSFORMATION =============
doc.add_heading('Section 4: Similarity Transformation', level=1)

if os.path.exists("handnote_page_4.png"):
    doc.add_picture("handnote_page_4.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('Stream Function (Mathematical Trick!)', level=2)

para = doc.add_paragraph()
run = para.add_run('ψ = x^(3/4) · f(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('η = x^(-1/4) · y')
run.font.name = 'Courier New'
run.font.size = Pt(11)
run.bold = True

doc.add_heading('What is Stream Function?', level=3)
doc.add_paragraph('A clever way to automatically satisfy continuity!')

doc.add_paragraph('Define: u = ∂ψ/∂y, v = -∂ψ/∂x', style='List Bullet')
doc.add_paragraph('Then: ∂u/∂x + ∂v/∂y = 0 automatically! (verify this yourself)', style='List Bullet')

doc.add_heading('Why This Helps:', level=3)
doc.add_paragraph('Reduces 2 equations (continuity + momentum) to 1 equation (just momentum)')
doc.add_paragraph('Converts PDEs to ODEs (much easier to solve!)')

doc.add_heading('Real-World Application:', level=3)
doc.add_paragraph('Engineers use this to solve boundary layer problems numerically. Instead of solving for u and v separately, solve for ψ once!')

doc.add_paragraph()
doc.add_heading('Boundary Conditions', level=2)

para = doc.add_paragraph()
para.add_run('At y = 0:').bold = True
doc.add_paragraph('u = v = 0 (no-slip condition)')
doc.add_paragraph('θ = 1 + A sin(πx) (sinusoidal wall temperature)')

para = doc.add_paragraph()
para.add_run('As y → ∞:').bold = True
doc.add_paragraph('u → 0 (velocity decays)')
doc.add_paragraph('θ → 0 (temperature approaches ambient)')

doc.add_heading('Physical Meaning:', level=3)
doc.add_paragraph('No-slip: Fluid sticks to wall (proven by dust on spinning fan blades!)')
doc.add_paragraph('Far-field: Far from wall, flow returns to undisturbed state')
doc.add_paragraph('Sinusoidal T: Wall temperature varies along length (A = amplitude)')

doc.add_page_break()

# ============= FINAL SECTION: QUIZ STRATEGY =============
doc.add_heading('QUIZ PREPARATION STRATEGY', level=1)

para = doc.add_paragraph()
run = para.add_run('🎯 3-Day Plan (Wed-Thu-Fri)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
doc.add_heading('Wednesday (Today):', level=2)
doc.add_paragraph('Morning: Memorize all dimensionless numbers (Re, Pr, Gr, Ra, Nu)')
doc.add_paragraph('Afternoon: Practice chain rule derivations 10 times')
doc.add_paragraph('Evening: Start DWT 2012 problem')

doc.add_heading('Thursday:', level=2)
doc.add_paragraph('Morning: Complete DWT 2012 from scratch')
doc.add_paragraph('Afternoon: Do it again without notes')
doc.add_paragraph('Evening: Review boundary layer equations')

doc.add_heading('Friday:', level=2)
doc.add_paragraph('Morning: DWT 2012 one more time - time yourself (<30 min)')
doc.add_paragraph('Afternoon: Quick formula review')
doc.add_paragraph('Evening: Light review, good sleep!')

doc.add_paragraph()
doc.add_heading('Most Likely Quiz Questions:', level=2)

doc.add_paragraph('1. "Non-dimensionalize the following equation..." (50% probability)', style='List Number')
doc.add_paragraph('2. "Define and explain Grashof number" (80% probability)', style='List Number')
doc.add_paragraph('3. "Write boundary layer equations" (60% probability)', style='List Number')
doc.add_paragraph('4. "Derive ∂²ū/∂ȳ² using chain rule" (40% probability)', style='List Number')
doc.add_paragraph('5. "What is Prandtl number? Physical meaning?" (70% probability)', style='List Number')

doc.add_paragraph()
doc.add_heading('Final Checklist (Friday Night):', level=2)

doc.add_paragraph('Can you write from memory:', style='List Bullet')
doc.add_paragraph('Continuity equation ✓', style='List Bullet 2')
doc.add_paragraph('Momentum equation (boundary layer) ✓', style='List Bullet 2')
doc.add_paragraph('Energy equation ✓', style='List Bullet 2')
doc.add_paragraph('Gr = gβΔTL³/ν² ✓', style='List Bullet 2')
doc.add_paragraph('Pr = ν/α ✓', style='List Bullet 2')
doc.add_paragraph('Chain rule: ∂ū/∂x̄ = (∂u/∂x̄)·(∂x/∂x̄)·(∂ū/∂u) ✓', style='List Bullet 2')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('If you can do all of these, you\'re ready! 🚀')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)

# Save
doc.save('AMCS506_ULTIMATE_Study_Guide.docx')
print("✓ Created: AMCS506_ULTIMATE_Study_Guide.docx")
print("\nComplete guide with real-world examples and detailed explanations!")

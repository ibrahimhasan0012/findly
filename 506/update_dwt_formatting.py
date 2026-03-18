from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('DWT_2012_Complete_Solution.docx')

# Find and update the Engineering Quantities section
# We'll recreate it with better formatting

# Remove old content and add new formatted version
# Since we can't easily edit existing content, we'll add a new page with proper formatting

doc.add_page_break()
doc.add_heading('STEP 7: Engineering Quantities (Formatted)', level=1)

para = doc.add_paragraph()
run = para.add_run('Skin Friction Coefficient:')
run.bold = True
run.font.size = Pt(12)

# Add formatted equation
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# First part: (1/2)Gr_x^(1/4) C_f
run1 = para.add_run('(1/2) Gr')
run1.font.name = 'Cambria Math'
run1.font.size = Pt(12)

# Subscript x
run2 = para.add_run('x')
run2.font.subscript = True
run2.font.name = 'Cambria Math'
run2.font.size = Pt(10)

# Superscript 1/4
run3 = para.add_run('1/4')
run3.font.superscript = True
run3.font.name = 'Cambria Math'
run3.font.size = Pt(10)

# C_f
run4 = para.add_run(' C')
run4.font.name = 'Cambria Math'
run4.font.size = Pt(12)

run5 = para.add_run('f')
run5.font.subscript = True
run5.font.name = 'Cambria Math'
run5.font.size = Pt(10)

# Equals sign
run6 = para.add_run(' = ')
run6.font.name = 'Cambria Math'
run6.font.size = Pt(12)

# (∂u/∂Y)|_{Y=0}
run7 = para.add_run('(∂u/∂Y)')
run7.font.name = 'Cambria Math'
run7.font.size = Pt(12)

run8 = para.add_run('|')
run8.font.name = 'Cambria Math'
run8.font.size = Pt(12)

run9 = para.add_run('Y=0')
run9.font.subscript = True
run9.font.name = 'Cambria Math'
run9.font.size = Pt(10)

# Equals f''(x,0)
run10 = para.add_run(' = f\'\'(x,0)')
run10.font.name = 'Cambria Math'
run10.font.size = Pt(12)

doc.add_paragraph()

# Nusselt Number
para = doc.add_paragraph()
run = para.add_run('Nusselt Number (Heat Transfer Rate):')
run.bold = True
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Nu/Gr_x^(1/4)
run1 = para.add_run('Nu / Gr')
run1.font.name = 'Cambria Math'
run1.font.size = Pt(12)

run2 = para.add_run('x')
run2.font.subscript = True
run2.font.name = 'Cambria Math'
run2.font.size = Pt(10)

run3 = para.add_run('1/4')
run3.font.superscript = True
run3.font.name = 'Cambria Math'
run3.font.size = Pt(10)

# Equals
run4 = para.add_run(' = ')
run4.font.name = 'Cambria Math'
run4.font.size = Pt(12)

# -(∂θ/∂Y)|_{Y=0}
run5 = para.add_run('-(∂θ/∂Y)')
run5.font.name = 'Cambria Math'
run5.font.size = Pt(12)

run6 = para.add_run('|')
run6.font.name = 'Cambria Math'
run6.font.size = Pt(12)

run7 = para.add_run('Y=0')
run7.font.subscript = True
run7.font.name = 'Cambria Math'
run7.font.size = Pt(10)

# Equals -θ'(x,0)
run8 = para.add_run(' = -θ\'(x,0)')
run8.font.name = 'Cambria Math'
run8.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Add explanation
para = doc.add_paragraph()
run = para.add_run('Physical Meaning:')
run.bold = True

doc.add_paragraph('C_f: Dimensionless shear stress at the wall (skin friction)', style='List Bullet')
doc.add_paragraph('Nu: Dimensionless heat transfer rate at the wall', style='List Bullet')
doc.add_paragraph('f\'\'(x,0): Second derivative of stream function at wall', style='List Bullet')
doc.add_paragraph('θ\'(x,0): Temperature gradient at wall (negative for heat transfer from wall)', style='List Bullet')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Note: ')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
para.add_run('These match equation (26) in the DWT 2012 paper exactly!')

# Save
doc.save('DWT_2012_Complete_Solution.docx')
print("✓ Updated: DWT_2012_Complete_Solution.docx")
print("\nAdded properly formatted engineering quantities!")

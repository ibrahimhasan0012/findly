from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DWT 2012 Problem - Complete Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.add_run("Natural Convection from Vertical Plate with Heat Source in Stratified Medium").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# Problem Statement
doc.add_heading('Problem Statement', level=1)

doc.add_paragraph('Natural convection flow from an isothermal vertical plate with uniform heat source embedded in a stratified medium.')

doc.add_paragraph()
doc.add_heading('Given:', level=2)
doc.add_paragraph('Wall temperature: Tᵥᵥ = constant')
doc.add_paragraph('Ambient temperature: T∞(x) = T₀ + B(x/L) (stratified)')
doc.add_paragraph('Internal heat generation: Q(T - T∞)')
doc.add_paragraph('Gravity: g (downward)')

doc.add_page_break()

# STEP 1
doc.add_heading('STEP 1: Dimensional Governing Equations', level=1)

para = doc.add_paragraph()
para.add_run('Continuity:').bold = True
para = doc.add_paragraph()
para.add_run('∂ū/∂x̄ + ∂v̄/∂ȳ = 0').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Momentum (boundary layer approximation):').bold = True
para = doc.add_paragraph()
para.add_run('ū(∂ū/∂x̄) + v̄(∂ū/∂ȳ) = ν(∂²ū/∂ȳ²) + gβ(T - T∞,x)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Energy:').bold = True
para = doc.add_paragraph()
para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²) + Q(T - T∞,x)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Boundary Conditions:').bold = True
para = doc.add_paragraph()
para.add_run('At ȳ = 0: ū = v̄ = 0, T = Tᵥᵥ').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('As ȳ → ∞: ū → 0, T → T∞,x').font.name = 'Courier New'

doc.add_page_break()

# STEP 2
doc.add_heading('STEP 2: Define Reference Scales', level=1)

doc.add_paragraph('Length scale: L (reference length)')
doc.add_paragraph('Temperature scale: ΔT = Tᵥᵥ - T₀')
doc.add_paragraph('Velocity scale: U = (ν/L)·Gr^(1/2)')

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Grashof Number:').bold = True
para = doc.add_paragraph()
run = para.add_run('Gr = gβ(Tᵥᵥ - T₀)L³/ν²')
run.font.name = 'Courier New'
run.bold = True

doc.add_page_break()

# STEP 3
doc.add_heading('STEP 3: Define Dimensionless Variables', level=1)

para = doc.add_paragraph()
para.add_run('Coordinates:').bold = True
para = doc.add_paragraph()
para.add_run('x = x̄/L').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('y = (Gr^(1/4)/L)·ȳ  (stretched coordinate!)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Velocities:').bold = True
para = doc.add_paragraph()
para.add_run('u = (L/ν)·Gr^(-1/2)·ū').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('v = (L/ν)·Gr^(-1/4)·v̄').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Temperature:').bold = True
para = doc.add_paragraph()
para.add_run('θ = (T - T∞,x)/(Tᵥᵥ - T₀)').font.name = 'Courier New'

doc.add_page_break()

# STEP 4
doc.add_heading('STEP 4: Dimensionless Equations', level=1)

para = doc.add_paragraph()
run = para.add_run('After substitution and simplification:')
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Continuity:').bold = True
para = doc.add_paragraph()
run = para.add_run('∂u/∂x + ∂v/∂y = 0')
run.font.name = 'Courier New'
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Momentum:').bold = True
para = doc.add_paragraph()
run = para.add_run('u(∂u/∂x) + v(∂u/∂y) = ∂²u/∂y² + θ')
run.font.name = 'Courier New'
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Energy:').bold = True
para = doc.add_paragraph()
run = para.add_run('u(∂θ/∂x) + v(∂θ/∂y) = (1/Pr)(∂²θ/∂y²) + λθ')
run.font.name = 'Courier New'
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Parameters:').bold = True
para = doc.add_paragraph()
para.add_run('Pr = ν/α (Prandtl number)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('λ = QL²/ν (heat generation parameter)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('S = B/(Tᵥᵥ - T₀)·Gr^(1/2) (stratification parameter)').font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Boundary Conditions:').bold = True
para = doc.add_paragraph()
para.add_run('At y = 0: u = v = 0, θ = 1 - Sx').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('As y → ∞: u → 0, θ → 0').font.name = 'Courier New'

doc.add_page_break()

# STEP 5
doc.add_heading('STEP 5: Similarity Transformation', level=1)

para = doc.add_paragraph()
para.add_run('Introduce stream function and similarity variable:').bold = True

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('ψ = x^(3/4)·f(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(12)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('η = x^(-1/4)·y')
run.font.name = 'Courier New'
run.font.size = Pt(12)
run.bold = True

para = doc.add_paragraph()
run = para.add_run('θ = θ(x,η)')
run.font.name = 'Courier New'
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Velocity components:').bold = True
para = doc.add_paragraph()
para.add_run('u = ∂ψ/∂y = x^(1/2)·f\'(x,η)').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('v = -∂ψ/∂x = (1/4)x^(-1/4)·[ηf\' - 3f] - x^(3/4)·∂f/∂x').font.name = 'Courier New'

doc.add_page_break()

# STEP 6
doc.add_heading('STEP 6: Final ODE System', level=1)

para = doc.add_paragraph()
run = para.add_run('After substituting similarity variables into PDEs:')
run.bold = True

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Momentum ODE:').bold = True
para = doc.add_paragraph()
run = para.add_run("f''' + (3/4)f·f'' - (1/2)(f')² + θ = x(f'·∂f'/∂x - f''·∂f/∂x)")
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Energy ODE:').bold = True
para = doc.add_paragraph()
run = para.add_run("(1/Pr)θ'' + (3/4)f·θ' + λθ = x(f'·∂θ/∂x - θ'·∂f/∂x)")
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 200)

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Boundary Conditions:').bold = True
para = doc.add_paragraph()
para.add_run('At η = 0: f = f\' = 0, θ = 1 - Sx').font.name = 'Courier New'
para = doc.add_paragraph()
para.add_run('As η → ∞: f\' → 0, θ → 0').font.name = 'Courier New'

doc.add_page_break()

# STEP 7
doc.add_heading('STEP 7: Engineering Quantities', level=1)

para = doc.add_paragraph()
para.add_run('Skin Friction Coefficient:').bold = True
para = doc.add_paragraph()
run = para.add_run("(1/2)Gr^(1/4)·Cf = (∂u/∂y)|_{y=0} = f''(x,0)")
run.font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Nusselt Number (Heat Transfer Rate):').bold = True
para = doc.add_paragraph()
run = para.add_run("Nu/Gr^(1/4) = -(∂θ/∂y)|_{y=0} = -θ'(x,0)")
run.font.name = 'Courier New'

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('These are what you calculate and plot!')
run.italic = True
run.font.color.rgb = RGBColor(200, 0, 0)

doc.add_page_break()

# Summary
doc.add_heading('SUMMARY: Complete Solution Process', level=1)

doc.add_paragraph('1. Start with dimensional equations (continuity, momentum, energy)', style='List Number')
doc.add_paragraph('2. Define reference scales (L, ΔT, U based on Gr)', style='List Number')
doc.add_paragraph('3. Make variables dimensionless (x, y, u, v, θ)', style='List Number')
doc.add_paragraph('4. Substitute and simplify to get dimensionless PDEs', style='List Number')
doc.add_paragraph('5. Introduce similarity transformation (ψ, η)', style='List Number')
doc.add_paragraph('6. Convert PDEs to ODEs', style='List Number')
doc.add_paragraph('7. Solve ODEs numerically to get f(η) and θ(η)', style='List Number')
doc.add_paragraph('8. Calculate engineering quantities (Cf, Nu)', style='List Number')

doc.add_paragraph()
doc.add_heading('Key Parameters:', level=2)
doc.add_paragraph('Pr = 0.7 (air) or 7.0 (water)')
doc.add_paragraph('λ: heat generation parameter (λ > 0 for heating, λ < 0 for cooling)')
doc.add_paragraph('S: stratification parameter (how much ambient T varies with height)')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('This is EXACTLY what your handnotes pages 1-7 show!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)

# Save
doc.save('DWT_2012_Complete_Solution.docx')
print("✓ Created: DWT_2012_Complete_Solution.docx")
print("\nStandalone DWT 2012 solution document created!")

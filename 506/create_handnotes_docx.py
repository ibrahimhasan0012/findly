from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document for handwritten notes
doc = Document()

# Add title
title = doc.add_heading('AMCS 506 - Handwritten Notes from Feb 7 Class', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph()
info.add_run("Professor's handwritten notes from class on February 7, 2026").italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph("_" * 80)
doc.add_paragraph()

# Add each page as an image
for page_num in range(1, 8):
    img_file = f"handnote_page_{page_num}.png"
    
    if os.path.exists(img_file):
        # Add page header
        doc.add_heading(f'Page {page_num}', level=1)
        
        # Add image
        try:
            doc.add_picture(img_file, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Error loading image: {e}")
        
        # Add page break except for last page
        if page_num < 7:
            doc.add_page_break()

# Save document
doc.save('AMCS506_Handwritten_Notes_Feb7.docx')
print("Created: AMCS506_Handwritten_Notes_Feb7.docx")

# List all created files
print("\n" + "="*60)
print("All DOCX files created in AMCS 261/506 folder:")
print("="*60)
print("1. AMCS506_Quiz_Study_Guide.docx")
print("2. AMCS506_NonDimensionalization_Solution.docx")
print("3. AMCS506_Formula_Sheet.docx")
print("4. AMCS506_Handwritten_Notes_Feb7.docx")
print("="*60)

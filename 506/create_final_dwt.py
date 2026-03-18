from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DWT 2012 Problem - Complete Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 80)
doc.add_paragraph()

# Momentum equation
doc.add_heading('Momentum (boundary layer approximation):', level=2)
para = doc.add_paragraph()
run = para.add_run('ū(∂ū/∂x̄) + v̄(∂ū/∂ȳ) = ν(∂²ū/∂ȳ²) + gβ(T - T∞,x)')
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()

# Energy equation  
doc.add_heading('Energy:', level=2)
para = doc.add_paragraph()
run = para.add_run('ū(∂T/∂x̄) + v̄(∂T/∂ȳ) = α(∂²T/∂ȳ²) + Q(T - T∞,x)')
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()

# Boundary conditions
doc.add_heading('Boundary Conditions:', level=2)
para = doc.add_paragraph()
run = para.add_run('At ȳ = 0: ū = v̄ = 0, T = Tᵥᵥ')
run.font.name = 'Courier New'
run.font.size = Pt(11)

para = doc.add_paragraph()
run = para.add_run('As ȳ → ∞: ū → 0, T → T∞,x')
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_page_break()

# Engineering quantities - EXACTLY like equation (26)
doc.add_heading('Engineering Quantities (Equation 26):', level=1)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Skin Friction Coefficient:')
run.bold = True
run.font.size = Pt(12)

para = doc.add_paragraph()
run = para.add_run('(1/2)Grₓ¹ᐟ⁴Cₓ = (∂u/∂Y)|ᵧ₌₀ = f\'\'(x,0)')
run.font.name = 'Courier New'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Nusselt Number:')
run.bold = True
run.font.size = Pt(12)

para = doc.add_paragraph()
run = para.add_run('Nu/Grₓ¹ᐟ⁴ = -(∂θ/∂Y)|ᵧ₌₀ = -θ\'(x,0)')
run.font.name = 'Courier New'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Alternative plain ASCII version
para = doc.add_paragraph()
run = para.add_run('Plain ASCII version (if symbols don\'t display):')
run.bold = True
run.italic = True

para = doc.add_paragraph()
run = para.add_run('(1/2)Gr_x^(1/4) C_f = (∂u/∂Y)|_{Y=0} = f\'\'(x,0)')
run.font.name = 'Courier New'
run.font.size = Pt(11)

para = doc.add_paragraph()
run = para.add_run('Nu/Gr_x^(1/4) = -(∂θ/∂Y)|_{Y=0} = -θ\'(x,0)')
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('✓ These match equation (26) in the DWT 2012 paper')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

# Save
doc.save('DWT_2012_Solution_FINAL.docx')
print("✓ Created: DWT_2012_Solution_FINAL.docx")
print("\nNew file with both Unicode and ASCII versions!")
